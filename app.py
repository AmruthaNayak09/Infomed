from flask import Flask, request, render_template, send_from_directory, send_file, jsonify
import csv, os
from datetime import datetime
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches

app = Flask(__name__)
DATA_FILE = "patient_data.csv"

def set_detected_user(name):
    global detected_user
    detected_user = name

def generate_patient_id():
    today = datetime.now().strftime("%Y%m%d")
    count = len(load_patient_data())
    return f"{today}{count+1:03d}"

def load_patient_data():
    patients = {}
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, newline='') as f:
            reader = csv.DictReader(f)
            for row in reader:
                patients[row["name"]] = row
    return patients

def save_patient_data(patients):
    with open(DATA_FILE, "w", newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=[
            "id", "name", "dob", "age", "gender", "weight", "height", "bmi",
            "blood_pressure", "heart_rate", "blood", "blood_glucose", "allergies",
            "med_history", "email", "phone", "visit_date",
            "past_visits", "prescriptions", "image"
        ])
        writer.writeheader()
        for p in patients.values():
            writer.writerow(p)

def calculate_age(dob):
    try:
        dob_dt = datetime.strptime(dob, "%Y-%m-%d")
        today = datetime.today()
        return today.year - dob_dt.year - ((today.month, today.day) < (dob_dt.month, dob_dt.day))
    except:
        return ""

def calculate_bmi(weight, height):
    try:
        weight = float(weight)
        height = float(height) / 100
        return round(weight / (height ** 2), 2)
    except:
        return ""

@app.template_filter('datetimeformat')
def datetimeformat(value):
    return datetime.fromtimestamp(value).strftime('%Y-%m-%d')

@app.route("/", methods=["GET", "POST"])
def portal():
    global detected_user
    name_query = request.args.get("name")
    if name_query:
        detected_user = name_query

    if detected_user is None:
        return "No face detected yet. Please wait...", 200

    name = detected_user
    patients = load_patient_data()

    if name and name in patients:
        patient = patients[name]
    else:
        patient = {
            "id": generate_patient_id(),
            "name": name,
            "dob": "", "age": "", "gender": "", "weight": "", "height": "",
            "bmi": "", "blood_pressure": "", "heart_rate": "", "blood": "", "blood_glucose": "",
            "allergies": "", "med_history": "", "email": "", "phone": "",
            "visit_date": datetime.now().strftime("%Y-%m-%d"),
            "past_visits": "", "prescriptions": "",
            "image": f"/faces/{name}/front.jpg" if name else ""
        }

    if request.method == "POST":
        form = request.form
        dob = form.get("dob", "")
        weight = form.get("weight", "")
        height = form.get("height", "")
        prescriptions = []
        for key in form:
            if key.startswith("tablet_"):
                index = key.split("_")[1]
                tablet = form.get(f"tablet_{index}")
                duration = form.get(f"duration_{index}", "")
                food = form.get(f"food_{index}", "")
                dosage = form.get(f"time_{index}", "")  # now using time_ as 1-0-1
                if tablet:
                    prescriptions.append(f"{tablet} - {duration} days - {food} - {dosage}")

        patient.update({
            "dob": dob,
            "name": name,
            "age": form.get("age") or calculate_age(dob),
            "gender": form.get("gender", ""),
            "weight": weight,
            "height": height,
            "bmi": calculate_bmi(weight, height),
            "blood_pressure": form.get("blood_pressure", ""),
            "heart_rate": form.get("heart_rate", ""),
            "blood": form.get("blood", ""),
            "blood_glucose": form.get("blood_glucose", ""),
            "allergies": form.get("allergies", ""),
            "med_history": form.get("med_history", ""),
            "email": form.get("email", ""),
            "phone": form.get("phone", ""),
            "visit_date": form.get("visit_date", datetime.now().strftime("%Y-%m-%d")),
            "past_visits": ",".join([k for k in ["past_yes", "past_no"] if form.get(k)]),
            "prescriptions": "|".join(prescriptions),
            "image": f"/faces/{name}/front.jpg"
        })

        patients[name] = patient
        save_patient_data(patients)
        return jsonify({"success": True})

    return render_template("doctor_portal.html", data=patient, timestamp=datetime.now().timestamp())

@app.route("/download")
def download():
    if detected_user is None:
        return "No patient selected", 400

    patients = load_patient_data()
    patient = patients.get(detected_user)
    if not patient:
        return "Patient data not found", 404

    doc = Document()
    doc.add_heading("Patient Report", 0)

    image_path = os.path.join("known_faces", patient["name"], "front.jpg")
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(1.5))

    for key, value in patient.items():
        if key != "image":
            doc.add_paragraph(f"{key.capitalize()}: {value}")

    filename = f"{patient['name'].replace(' ', '_')}_report.docx"
    doc.save(filename)
    return send_file(filename, as_attachment=True)

@app.route("/get_patient")
def get_patient():
    name = request.args.get("name")
    patients = load_patient_data()
    return jsonify(patients.get(name, {}))

@app.route("/faces/<name>/<filename>")
def serve_face(name, filename):
    path = os.path.join("known_faces", name)
    return send_from_directory(path, filename)
